#!/usr/bin/python3

import docx
import sys
import openpyxl
import os
import wget
import re



from openpyxl.styles import Alignment
from openpyxl.styles import PatternFill
from zipfile import ZipFile
from win32com import client as wc

group = {"Radio Network Layer cause":"Radio Network Layer", "Transport Layer cause":"Transport Layer", "Protocol cause":"Protocol", "Miscellaneous cause":"Misc", "Transport Network Layer cause":"Transport Layer", "NAS cause":"NAS"}
excel_sheet_name="Cause Enumerations"

protocol = {}


protocol_col = 'A'
group_col = 'B'
code_col = 'C'
descripiton_col = 'D'

cause_pair = []
rrc_est_pair = []

current_protocol = ""

def save_cause(tbl_name, tbl):
    for idx in range(1, len(tbl.rows)):
        cause_pair.append((current_protocol, group[tbl_name], tbl.cell(idx,0).paragraphs[0].text))

def find_cause_tbl(file_name):
    print("\nParsing cause value from " + os.getcwd() + "\\" + file_name)
    doc=docx.Document(os.getcwd() + "\\" + file_name)
    for tbl in doc.tables:
        try:
            cell = tbl.cell(0,0)
            if tbl.cell(0,0).paragraphs[0].text in group:
                save_cause(tbl.cell(0,0).paragraphs[0].text, tbl)
        except:
            print("parse {} exception".format(tbl))
    if "38413" in file_name: #ngap there is rrc cause
        for rrc_pair in rrc_est_pair:
            cause_pair.append((rrc_pair[0], rrc_pair[1], rrc_pair[2]))

def fill_misc(sheet):
    #title
    sheet['a1'].fill = PatternFill(start_color='92D050', end_color='92D050', fill_type="solid")
    sheet['a1'].value = "Protocol Type"
    sheet['a1'].alignment = Alignment(vertical='center',horizontal='center')

    sheet['b1'].fill = PatternFill(start_color='92D050', end_color='92D050', fill_type="solid")
    sheet['b1'].value = "Cause Group"
    sheet['b1'].alignment = Alignment(vertical='center',horizontal='center')

    sheet['c1'].fill = PatternFill(start_color='92D050', end_color='92D050', fill_type="solid")
    sheet['c1'].value = "Cause Code"
    sheet['c1'].alignment = Alignment(vertical='center',horizontal='center')

    sheet['d1'].fill = PatternFill(start_color='92D050', end_color='92D050', fill_type="solid")
    sheet['d1'].value = "Description"
    sheet['d1'].alignment = Alignment(vertical='center',horizontal='center')

    sheet['f1'].value = "Auto generated, Do not Edit."

total_idx = 1 #excel col start with 1
def put_to_excel(new_sheet):
    global total_idx
    group = "dump"
    group_start_idx = group_end_idx = total_idx
    protocol = "dump"
    protocol_start_idx = protocol_end_idx = total_idx

    for pair in cause_pair:
        if protocol != pair[0]: #new protocol
            cell_idx = protocol_col + str(total_idx + 1)
            new_sheet[cell_idx].value = pair[0]
            protocol = pair[0]
            if protocol_start_idx != protocol_end_idx:
                new_sheet.merge_cells(protocol_col + str(protocol_start_idx + 1) + ":" + protocol_col + str(protocol_end_idx))
                new_sheet[protocol_col + str(protocol_start_idx + 1)].alignment = Alignment(vertical='center')
                #print("merge " + protocol_col + str(protocol_start_idx + 1) + ":" + protocol_col + str(protocol_end_idx))
            protocol_start_idx = total_idx

        if group != pair[1]: # new group
            code_num = 0
            cell_idx = group_col + str(total_idx + 1)
            new_sheet[cell_idx].value = pair[1]
            group = pair[1]
            if group_start_idx != group_end_idx:
                new_sheet.merge_cells(group_col + str(group_start_idx + 1) + ":" + group_col + str(group_end_idx))
                new_sheet[group_col + str(group_start_idx + 1)].alignment = Alignment(vertical='center')
                #print("merge " + group_col + str(group_start_idx + 1) + ":" + group_col + str(group_end_idx))
            group_start_idx = total_idx

        cell_idx = code_col + str(total_idx + 1) #first line is title
        new_sheet[cell_idx].value = code_num
        cell_idx = descripiton_col + str(total_idx + 1)
        new_sheet[cell_idx].value = pair[2]
        code_num += 1
        total_idx += 1
        protocol_end_idx = total_idx
        group_end_idx = total_idx

    if protocol_start_idx != protocol_end_idx:
        new_sheet.merge_cells(protocol_col + str(protocol_start_idx + 1) + ":" + protocol_col + str(protocol_end_idx))
        new_sheet[protocol_col + str(protocol_start_idx + 1)].alignment = Alignment(vertical='center')
        #print("merge " + protocol_col + str(protocol_start_idx + 1) + ":" + protocol_col + str(protocol_end_idx))
    if group_start_idx != group_end_idx:
        new_sheet.merge_cells(group_col + str(group_start_idx + 1) + ":" + group_col + str(group_end_idx))
        new_sheet[group_col + str(group_start_idx + 1)].alignment = Alignment(vertical='center')
        #print("merge " + group_col + str(group_start_idx + 1) + ":" + group_col + str(group_end_idx))

def doc_to_docx(file_name):
    try:
        word = wc.Dispatch("Word.Application")
        doc = word.Documents.Open(os.getcwd() + "\\" + file_name)
        doc.SaveAs("{}x".format(os.getcwd() + "\\" + file_name), 12)
    except:
        print("Can not convert {}} to docx.".format(file_name))
    finally:
        doc.Close()

def download_spec(url):
    #TODO auto generate file names
    try:
        response = wget.download(url)
    except:
        print("Download {} specs failed.".format(url))

def parse_rrc(file_name):
    print("\nParsing cause value from " + os.getcwd() + "\\" + file_name)
    doc=docx.Document(os.getcwd() + "\\" + file_name)
    all_para=""
    for context in doc.paragraphs:
        all_para += context.text
    for cause in re.search("rlf-Cause-r16.*?}",all_para, re.MULTILINE).group().split("{")[-1].split(","):
        cause_pair.append((current_protocol, "RRC RLF", cause.strip().removesuffix("}"))) 
    for cause in re.search("EstablishmentCause ::=.*?}",all_para, re.MULTILINE).group().split("{")[-1].split(","):
        cause_pair.append((current_protocol, "RRC Establishment", cause.strip().removesuffix("}")))
        rrc_est_pair.append(("NGAP", "RRC Establishment", cause.strip().removesuffix("}")))#for appending to ngap part 
    for cause in re.search("ReestablishmentCause ::=.*?}",all_para, re.MULTILINE).group().split("{")[-1].split(","):
        cause_pair.append((current_protocol, "RRC Reestablishment", cause.strip().removesuffix("}")))
    for cause in re.search("ResumeCause ::=.*?}",all_para, re.MULTILINE).group().split("{")[-1].split(","):
        cause_pair.append((current_protocol, "RRC Resume", cause.strip().removesuffix("}")))
    for cause in re.search("FailureReportSCG ::=.*?}",all_para, re.MULTILINE).group().split("{")[-1].split(","):
        cause_pair.append((current_protocol, "ScgFailureType", cause.strip().removesuffix("}")))
    for cause in re.search("failureType-v1610.*?}",all_para, re.MULTILINE).group().split("{")[-1].split(","):
        cause_pair.append((current_protocol, "ScgFailureType-v1610", cause.strip().removesuffix("}")))

def main():
    global current_protocol
    global protocol
    if len(sys.argv) < 3:
        print("Usage: cause_distribute.py <excel file> <3GPP_release_version> [S1AP]")
        exit(1)
    protocol = { 
           ("RRC", "https://www.3gpp.org/ftp/Specs/archive/38_series/38.331/38331-"+sys.argv[2]+".zip"),
           ("XnAP","https://www.3gpp.org/ftp/Specs/archive/38_series/38.423/38423-"+sys.argv[2]+".zip"),
           ("E1AP","https://www.3gpp.org/ftp/Specs/archive/37_series/37.483/37483-"+sys.argv[2]+".zip"),
           ("F1AP","https://www.3gpp.org/ftp/Specs/archive/38_series/38.473/38473-"+sys.argv[2]+".zip"),
           ("X2AP","https://www.3gpp.org/ftp/Specs/archive/36_series/36.423/36423-"+sys.argv[2]+".zip"),
           ("NGAP","https://www.3gpp.org/ftp/Specs/archive/38_series/38.413/38413-"+sys.argv[2]+".zip")
    }

    if len(sys.argv) == 4 and sys.argv[3].upper() == "S1AP":
        protocol.add(("S1AP","https://www.3gpp.org/ftp/Specs/archive/36_series/36.413/36413-"+sys.argv[2]+".zip"))
        print("S1AP cause code enabled")
    else:
        if len(sys.argv) >= 4:
            print("Usage: cause_distribute.py <excel file> <3GPP_release_version> [S1AP]")
            exit(1)

    wb = openpyxl.load_workbook(sys.argv[1])
    new_sheet = wb.create_sheet(title = excel_sheet_name + "1")
    old_sheet = wb[excel_sheet_name]
 
    for p in sorted(protocol, reverse=True):
        spec_name = p[1].split("/")[-1]
        if not os.path.exists(spec_name):
            download_spec(p[1])
        with ZipFile(spec_name, 'r') as zipObj:
            zipObj.extractall()
        current_protocol = p[0]
        cause_pair.clear()
        if not os.path.exists(spec_name.replace("zip","docx")):
            doc_to_docx(spec_name.replace("zip","doc"))
        if p[0] == "RRC":
            parse_rrc(spec_name.replace("zip","docx"))
        else:
            find_cause_tbl(spec_name.replace("zip","docx"))
        put_to_excel(new_sheet)
    #save excel
    fill_misc(new_sheet)
    wb.remove(old_sheet)
    new_sheet.title = excel_sheet_name
    wb.save(sys.argv[1])
    wb.close()

main()
